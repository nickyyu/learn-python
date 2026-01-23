import sys
from dataclasses import dataclass, field
from datetime import datetime
from typing import List

from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget,
    QVBoxLayout, QHBoxLayout, QLabel, QLineEdit,
    QTextEdit, QPushButton, QListWidget, QRadioButton,
    QButtonGroup, QMessageBox, QTableWidget, QTableWidgetItem,
    QFileDialog, QDateTimeEdit
)
from docx import Document


# ======================
# 数据模型
# ======================

@dataclass
class AppDeploy:
    app_name: str = ""
    deploy_type: str = "vm"   # vm / container
    build_version: str = ""
    ip_nodes: List[str] = field(default_factory=list)
    image: str = ""
    namespace: str = ""


@dataclass
class DatabaseChange:
    name: str
    change_type: str
    remark: str


@dataclass
class CronJob:
    name: str
    schedule: str
    remark: str


# ======================
# Word 生成器
# ======================

class WordGenerator:
    def generate(self, basic, apps, dbs, configs, crons, path):
        doc = Document()
        doc.add_heading("上线方案", level=1)

        doc.add_heading("一、基本信息", level=2)
        doc.add_paragraph(f"需求内容：{basic['requirement']}")
        doc.add_paragraph(f"上线时间：{basic['time']}")
        doc.add_paragraph(f"上线版本：{basic['version']}")
        doc.add_paragraph(f"关联系统：{basic['systems']}")

        doc.add_heading("二、发布应用", level=2)
        for i, app in enumerate(apps, 1):
            doc.add_heading(f"{i}. {app.app_name}", level=3)
            doc.add_paragraph(f"部署方式：{'虚机' if app.deploy_type == 'vm' else '容器'}")
            doc.add_paragraph(f"构建版本：{app.build_version}")

            if app.deploy_type == "vm":
                doc.add_paragraph("IP 节点：")
                for ip in app.ip_nodes:
                    doc.add_paragraph(ip, style="List Bullet")
            else:
                doc.add_paragraph(f"镜像：{app.image}")
                doc.add_paragraph(f"命名空间：{app.namespace}")

        doc.add_heading("三、数据库变更", level=2)
        for db in dbs:
            doc.add_paragraph(f"{db.name} | {db.change_type} | {db.remark}")

        doc.add_heading("四、配置变更", level=2)
        for c in configs:
            doc.add_paragraph(c, style="List Bullet")

        doc.add_heading("五、定时任务", level=2)
        for c in crons:
            doc.add_paragraph(f"{c.name} | {c.schedule} | {c.remark}")

        doc.save(path)


# ======================
# 主窗口
# ======================

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("上线方案自动生成工具")
        self.resize(1000, 700)

        self.apps: List[AppDeploy] = []

        self.tabs = QTabWidget()
        self.setCentralWidget(self.tabs)

        self.init_basic_tab()
        self.init_app_tab()
        self.init_other_tabs()
        self.init_bottom_bar()

    # ---------- 基本信息 ----------
    def init_basic_tab(self):
        tab = QWidget()
        layout = QVBoxLayout(tab)

        self.req_input = QTextEdit()
        self.time_input = QDateTimeEdit(datetime.now())
        self.version_input = QLineEdit()
        self.system_input = QTextEdit()

        layout.addWidget(QLabel("需求内容"))
        layout.addWidget(self.req_input)
        layout.addWidget(QLabel("上线时间"))
        layout.addWidget(self.time_input)
        layout.addWidget(QLabel("上线版本"))
        layout.addWidget(self.version_input)
        layout.addWidget(QLabel("关联系统"))
        layout.addWidget(self.system_input)

        self.tabs.addTab(tab, "基本信息")

    # ---------- 发布应用 ----------
    def init_app_tab(self):
        tab = QWidget()
        main = QHBoxLayout(tab)

        # 左侧列表
        left = QVBoxLayout()
        self.app_list = QListWidget()
        btn_add = QPushButton("➕ 新增应用")
        btn_del = QPushButton("➖ 删除应用")

        left.addWidget(self.app_list)
        left.addWidget(btn_add)
        left.addWidget(btn_del)

        # 右侧表单
        right = QVBoxLayout()
        self.app_name = QLineEdit()
        self.build_version = QLineEdit()

        self.rb_vm = QRadioButton("虚机")
        self.rb_container = QRadioButton("容器")
        self.rb_vm.setChecked(True)

        self.ip_text = QTextEdit()
        self.image_input = QLineEdit()
        self.ns_input = QLineEdit()

        right.addWidget(QLabel("应用名称"))
        right.addWidget(self.app_name)
        right.addWidget(QLabel("构建版本"))
        right.addWidget(self.build_version)
        right.addWidget(self.rb_vm)
        right.addWidget(self.rb_container)
        right.addWidget(QLabel("IP 节点 / 容器信息"))
        right.addWidget(self.ip_text)
        right.addWidget(QLabel("镜像（容器）"))
        right.addWidget(self.image_input)
        right.addWidget(QLabel("命名空间（容器）"))
        right.addWidget(self.ns_input)

        main.addLayout(left, 2)
        main.addLayout(right, 5)

        btn_add.clicked.connect(self.add_app)
        btn_del.clicked.connect(self.delete_app)
        self.app_list.currentRowChanged.connect(self.switch_app)

        self.tabs.addTab(tab, "发布应用")

    def add_app(self):
        app = AppDeploy(app_name=f"应用{len(self.apps)+1}")
        self.apps.append(app)
        self.app_list.addItem(app.app_name)
        self.app_list.setCurrentRow(len(self.apps) - 1)

    def delete_app(self):
        row = self.app_list.currentRow()
        if row < 0:
            return
        self.apps.pop(row)
        self.app_list.takeItem(row)
        if self.apps:
            self.app_list.setCurrentRow(0)
        else:
            self.clear_app_form()

    def switch_app(self, row):
        self.save_current_app()
        if row >= 0:
            self.load_app(row)

    def save_current_app(self):
        row = self.app_list.currentRow()
        if row < 0:
            return
        app = self.apps[row]
        app.app_name = self.app_name.text()
        app.build_version = self.build_version.text()
        app.deploy_type = "vm" if self.rb_vm.isChecked() else "container"
        app.ip_nodes = self.ip_text.toPlainText().splitlines()
        app.image = self.image_input.text()
        app.namespace = self.ns_input.text()
        self.app_list.item(row).setText(app.app_name)

    def load_app(self, row):
        app = self.apps[row]
        self.app_name.setText(app.app_name)
        self.build_version.setText(app.build_version)
        self.rb_vm.setChecked(app.deploy_type == "vm")
        self.rb_container.setChecked(app.deploy_type == "container")
        self.ip_text.setPlainText("\n".join(app.ip_nodes))
        self.image_input.setText(app.image)
        self.ns_input.setText(app.namespace)

    def clear_app_form(self):
        self.app_name.clear()
        self.build_version.clear()
        self.ip_text.clear()
        self.image_input.clear()
        self.ns_input.clear()

    # ---------- 其他 Tab ----------
    def init_other_tabs(self):
        self.db_table = self.create_table_tab("数据库", ["名称", "类型", "备注"])
        self.cfg_table = self.create_table_tab("配置", ["配置项"])
        self.cron_table = self.create_table_tab("定时任务", ["名称", "时间", "备注"])

    def create_table_tab(self, name, headers):
        tab = QWidget()
        layout = QVBoxLayout(tab)
        table = QTableWidget(0, len(headers))
        table.setHorizontalHeaderLabels(headers)
        btn = QPushButton("新增行")
        btn.clicked.connect(lambda: table.insertRow(table.rowCount()))
        layout.addWidget(table)
        layout.addWidget(btn)
        self.tabs.addTab(tab, name)
        return table

    # ---------- 底部 ----------
    def init_bottom_bar(self):
        btn = QPushButton("生成 Word")
        btn.clicked.connect(self.generate_word)
        self.statusBar().addPermanentWidget(btn)

    def generate_word(self):
        self.save_current_app()

        path, _ = QFileDialog.getSaveFileName(
            self, "保存 Word", "上线方案.docx", "Word (*.docx)"
        )
        if not path:
            return

        basic = {
            "requirement": self.req_input.toPlainText(),
            "time": self.time_input.dateTime().toString("yyyy-MM-dd HH:mm"),
            "version": self.version_input.text(),
            "systems": self.system_input.toPlainText()
        }

        dbs = []
        for r in range(self.db_table.rowCount()):
            dbs.append(DatabaseChange(
                self.db_table.item(r, 0).text(),
                self.db_table.item(r, 1).text(),
                self.db_table.item(r, 2).text()
            ))

        configs = [
            self.cfg_table.item(r, 0).text()
            for r in range(self.cfg_table.rowCount())
        ]

        crons = []
        for r in range(self.cron_table.rowCount()):
            crons.append(CronJob(
                self.cron_table.item(r, 0).text(),
                self.cron_table.item(r, 1).text(),
                self.cron_table.item(r, 2).text()
            ))

        WordGenerator().generate(basic, self.apps, dbs, configs, crons, path)
        QMessageBox.information(self, "成功", "Word 生成完成")


# ======================
# 程序入口
# ======================

if __name__ == "__main__":
    app = QApplication(sys.argv)
    w = MainWindow()
    w.show()
    sys.exit(app.exec())
